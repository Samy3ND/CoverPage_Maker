import streamlit as st
from docx import Document
from datetime import datetime
from io import BytesIO
import os
import win32com.client as win32
import pythoncom
from docx.shared import Pt
import tempfile


#student list for maping

roll_number_to_name={

 '022BIM001':'Aabha Kumhal',
 
 '022BIM003':'Aarchi Palikhel',
 '022BIM004':'Aarohan Shakya',
 '022BIM005':'Aayush Ghimire',
 '022BIM006':'Abhilasha Adhikari',
 '022BIM007':'Aelish Maharjan',
 '022BIM010':'Anish Katwal',
 '022BIM012':'Anuskha Shakya',
 '022BIM013':'Aprama Pokhrel',
 '022BIM014':'Arya Jyoti Bajracharya',
 '022BIM015':'Aryan Man Singh Pradhan',
 '022BIM016':'Avishek Bista',
 '022BIM017':'Babita Khadka',
 '022BIM018':'',
 '022BIM019':'Biju Shrestha',
 '022BIM020':'Bimmi Shrestha',
 '022BIM022':'Davish Shrestha',
 '022BIM023':'Deepika Dangol',
 '022BIM024':'Drishya Dangol',
 '022BIM025':'Gracy Rai',
 '022BIM026':'Hrishav Karmacharya',
 '022BIM028':'Ishan Thapa Magar',
 '022BIM029':'Krishtina Ranjit',
 '022BIM030':'Kritan Man Shrestha',
 '022BIM031':'Lyrica Rana',
 '022BIM032':'Mehebika Rai',
 '022BIM033':'Nidhi Rauniyar',
 '022BIM067':'Shubham Yadav',
 '022BIM035':'Nirusha Chalise',
 '022BIM036'	:'Nischal Shrestha',
 '022BIM037':'Nishant Pokherel',
 '022BIM038':'Pragya Chalise',
 '022BIM039'	:'Prakriti Acharya',
 '022BIM040'	:' Prapti Bajracharya',
 '022BIM041'	:'Pratik Maharjan',
 '022BIM042'	:'Prince Panthi',
 '022BIM044'	:'Rishi Kumar Panday',
 '022BIM045'	:'Rishna Joshi',
 '022BIN046'	:'Sachistha Gurung',
 '022BIM047'	:'Sampurna Poudyal',
 '022BIM048'	:'Samyog K.C.',
 '022BIM049'	:'Saraswoti Kapali',
 '022BIM050'	:'Sarthak Rupakheti',
 '022BIM051' :'Sashank Shahi',
 '022BIM052'	:'Shreskar Bista',
 '022BIM053'	:'Shriya Shakya',
 '022BIM054' :'Shriyanshu Dhakal',
 '022BIM055' :'Srijana Khatri',
 '022BIM056' :'Stuti Karanjeet',
 '022BIM057'	:'Subin Malla',
 '022BIM058' :'Sudhanshu Yadav',
 '022BIM059'	:'Sugam Rana',
 '022BIM060'	:'Sujan Pokharel',
 '022BIM062'	:'Sushan Narayan Dangol',
 '022BIM063'	:'Tisa Manandhar',
 '022BIM064' :'Triza Kafle',
 '022BIM065'	:'Vijan Dharel',
 '022BIM066' : 'Aakanksha Lamsal'  
}

# Subject Teacher name 

subject_to_teacher = {
    "System Design & Development [IT 242]": "Er. Sanjay Kumar Yadav",
    "Python": "Mr Ramesh Shahi [IT 243]",
    "Artificial Intelligence [IT 288]": "Er Nischal Shrestha",
    "Information Security [IT 244]": "Er. Saroj Shahi",
}

template_path = r"C:\Temp\Font.docx"  #Font Template Path

def replace_placeholders(doc, name, roll_number, lab_report_number, subject, teacher):
    current_date = datetime.now().strftime("%d-%m-%Y")
    replacements = {
        "{Name}": name,
        "{RollNumber}": roll_number,
        "{LabReportNumber}": lab_report_number,
        "{Date}": current_date,
        "{Subject}": subject,
        "{Teacher}": teacher,
    }

    # Function to replace text in the paragraph
def replace_in_paragraph(paragraph, replacements):
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                print(f"Replacing: {run.text} -> {run.text.replace(key, value)}")  
                run.text = run.text.replace(key, value)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    return doc

def replace_placeholders(doc, name, roll_number, lab_report_number, subject, teacher):
    current_date = datetime.now().strftime("%d-%m-%Y")
    replacements = {
        "{Name}": name,
        "{RollNumber}": roll_number,
        "{LabReportNumber}": lab_report_number,
        "{Date}": current_date,
        "{Subject}": subject,
        "{Teacher}": teacher,
    }

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                full_text = "".join(run.text for run in paragraph.runs)
                updated_text = full_text.replace(key, value)
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = updated_text

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            full_text = "".join(run.text for run in paragraph.runs)
                            updated_text = full_text.replace(key, value)
                            for run in paragraph.runs:
                                run.text = ""
                            paragraph.runs[0].text = updated_text

    return doc

def convert_docx_to_pdf(doc, output_name):
    pythoncom.CoInitialize()

    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False  # 
        print("Word application initialized.") 

        # Save the DOCX file to a temporary location using tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_doc_file:
            temp_doc_path = temp_doc_file.name
            if doc is None:
                raise ValueError("The DOCX document is not valid.")
            
            doc.save(temp_doc_path)  
            print(f"Saved DOCX to: {temp_doc_path}")  

        # Check if the DOCX file exists
        if not os.path.exists(temp_doc_path):
            raise FileNotFoundError(f"File not found: {temp_doc_path}")

        # Convert the DOCX file to PDF
        pdf_file_path = temp_doc_path.replace(".docx", ".pdf")
        print(f"PDF will be saved to: {pdf_file_path}") 

        # Attempt to open the DOCX file and save as PDF
        document = word.Documents.Open(temp_doc_path)
        if document is None:
            raise Exception(f"Failed to open document: {temp_doc_path}")

        document.Activate()  
        print(f"Document opened: {temp_doc_path}") 

        # Save the document as PDF
        document.SaveAs(pdf_file_path, FileFormat=17) 
        print(f"PDF saved to: {pdf_file_path}") 
    except Exception as e:
        print(f"Error during DOCX to PDF conversion: {e}")  
        raise Exception(f"Error during DOCX to PDF conversion: {e}")
    finally:
        try:
            if document:
                document.Close()
            if word:
                word.Quit()  
            print("Document closed and Word application quit.") 
        except Exception as close_error:
            print(f"Error during cleanup: {close_error}") 

        pythoncom.CoUninitialize()

    # Read the PDF into memory
    with open(pdf_file_path, "rb") as pdf_file:
        pdf_bytes = BytesIO(pdf_file.read())
    os.remove(temp_doc_path)
    os.remove(pdf_file_path)

    return pdf_bytes

# Streamlit UI
st.title("Lab Report Cover Page Maker")
st.header("Only Applicable for 5th Sem of SXC")

roll_number_input = st.text_input("Roll Number:").strip() 

if roll_number_input:
    name = roll_number_to_name.get(roll_number_input, "")
    if name:
        st.text_input("Name:", value=name, disabled=True)  
    else:
        st.warning(f"No name found for Roll Number {roll_number_input}. Please verify the input.")

subject = st.selectbox(
    "Subject:",
    ["System Design & Development [IT 242]", "Python [IT 243]", "Artificial Intelligence [IT 288]", "Information Security [IT 244]"]
)

teacher = subject_to_teacher.get(subject, "Teacher not assigned")
st.text_input("Teacher:", value=teacher, disabled=True) 

lab_report_number = st.number_input("Lab Report Number:", min_value=1, step=1)

if st.button("Generate Cover Page"):
    if name and roll_number_input and lab_report_number:
        try:
            with st.spinner("Generating your cover page..."):
                doc = Document(template_path)

                updated_doc = replace_placeholders(doc, name, roll_number_input, str(lab_report_number), subject, teacher)

                pdf_bytes = convert_docx_to_pdf(updated_doc, "cover_page")
            
    
                st.success("Cover page generated successfully!")
                st.download_button(
                    label="Download Cover Page",
                    data=pdf_bytes,
                    file_name="Lab_Report_Cover.pdf",
                    mime="application/pdf",
                )
        except Exception as e:
            st.error(f"Error: {e}")
    else:
        st.error("Please fill out all the fields!")
